from flask import Flask, request, jsonify, session
from flask_cors import CORS
import os
from datetime import datetime, timedelta
from dotenv import load_dotenv
import requests
import json
import itertools # Import itertools for cycling through keys
import pikepdf
import io
import os
import tempfile
import shutil
from flask import send_file, make_response
from PIL import Image, ImageDraw
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import A4, LETTER, LEGAL, A3
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import zipfile
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from pdf2image import convert_from_bytes
import pytesseract
from docx import Document
from docx.shared import Inches
from flask import session




# Load environment variables first
load_dotenv()

# Now import supabase client
from supabase_client import create_supabase_client, handle_supabase_error

app = Flask(__name__, static_folder='../frontend', static_url_path='')
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key')

# Configure CORS to allow all origins for development
CORS(app,
     supports_credentials=True,
     origins=["*"],  # Allow all origins for now
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
     allow_headers=["*"],
     expose_headers=["*"])


if os.getenv('RENDER'):  # Render sets this environment variable automatically
    # Production settings (HTTPS)
    app.config['SESSION_COOKIE_SECURE'] = True
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
else:
    # Development settings (HTTP)
    app.config['SESSION_COOKIE_SECURE'] = False
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
# Common session settings
app.config['SESSION_COOKIE_NAME'] = 'oneclick_session'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)

app.config['SESSION_COOKIE_DOMAIN'] = None  # Explicitly set to None for localhost
app.config['SESSION_COOKIE_PATH'] = '/'
app.config['SESSION_REFRESH_EACH_REQUEST'] = True
# Session configuration - UPDATED


@app.before_request
def log_session_info():
    """Log session information for debugging"""
    if request.path.startswith('/api/'):
        print(f"=== SESSION DEBUG ===")
        print(f"Path: {request.path}")
        print(f"Session ID: {session.sid if hasattr(session, 'sid') else 'None'}")
        print(f"Session data: {dict(session)}")
        print(f"Cookies received: {request.cookies}")
        print("=====================")

# Add this right after your session configuration

# Initialize Supabase client
try:
    supabase = create_supabase_client()
    print("Supabase client initialized successfully!")
except Exception as e:
    print(f"Error initializing Supabase: {e}")
    supabase = None

# --- A4F API Key Rotation Setup ---
A4F_API_KEYS_STR = os.environ.get('A4F_API_KEYS', '')
A4F_API_KEYS = [key.strip() for key in A4F_API_KEYS_STR.split(',') if key.strip()]

if not A4F_API_KEYS:
    print("Warning: No A4F API keys found. Image generation and Chat will not work.")
    current_a4f_api_key_iterator = None
else:
    current_a4f_api_key_iterator = itertools.cycle(A4F_API_KEYS)
    current_a4f_api_key = next(current_a4f_api_key_iterator)
    print(f"Initialized A4F API keys. Starting with: {current_a4f_api_key[:5]}...")

# --- Picsart API Key Rotation Setup ---
PICSART_API_KEYS_STR = os.environ.get('PICSART_API_KEYS', '')
PICSART_API_KEYS = [key.strip() for key in PICSART_API_KEYS_STR.split(',') if key.strip()]

if not PICSART_API_KEYS:
    print("Warning: No Picsart API keys found. Image enhancement will not work.")
    current_picsart_api_key_iterator = None
else:
    current_picsart_api_key_iterator = itertools.cycle(PICSART_API_KEYS)
    current_picsart_api_key = next(current_picsart_api_key_iterator)
    print(f"Initialized Picsart API keys. Starting with: {current_picsart_api_key[:5]}...")


@app.after_request
def after_request(response):
    """Add CORS headers to every response"""
    response.headers.add('Access-Control-Allow-Credentials', 'true')
    response.headers.add('Access-Control-Allow-Headers', '*')
    response.headers.add('Access-Control-Allow-Methods', '*')
    response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
    return response

# Serve frontend files
@app.route('/')
def serve_index():
    return app.send_static_file('index.html')

@app.route('/<path:path>')
def serve_static(path):
    if os.path.exists(os.path.join(app.static_folder, path)):
        return app.send_static_file(path)
    else:
        return app.send_static_file('index.html')

# API Routes
@app.route('/api/health')
def health_check():
    return jsonify({
        "status": "healthy",
        "supabase": "connected" if supabase else "disconnected",
        "message": "Server is running"
    })

@app.route('/api/register', methods=['POST', 'OPTIONS'])
def register():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    if not supabase:
        return jsonify({
            "success": False,
            "message": "Server configuration error"
        }), 500

    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        full_name = data.get('fullName')

        if not all([email, password, full_name]):
            return jsonify({
                "success": False,
                "message": "All fields are required"
            }), 400

        print(f"Registering user: {email}")

        # Create user in Supabase Auth
        auth_response = supabase.auth.sign_up({
            "email": email,
            "password": password,
        })

        if auth_response.user:
            # Store user data
            user_data = {
                "id": auth_response.user.id,
                "email": email,
                "full_name": full_name,
                "created_at": datetime.utcnow().isoformat()
            }

            user_response = supabase.table('users').insert(user_data).execute()

            response = jsonify({
                "success": True,
                "message": "Registration successful",
                "user": {
                    "id": auth_response.user.id,
                    "email": auth_response.user.email,
                    "full_name": full_name
                }
            })
            response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
            response.headers.add('Access-Control-Allow-Credentials', 'true')
            return response, 201

    except Exception as e:
        error_message = handle_supabase_error(e)
        response = jsonify({
            "success": False,
            "message": error_message
        })
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response, 500

@app.route('/api/login', methods=['POST', 'OPTIONS'])
def login():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    if not supabase:
        return jsonify({
            "success": False,
            "message": "Server configuration error"
        }), 500

    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return jsonify({
                "success": False,
                "message": "Email and password are required"
            }), 400

        print(f"Logging in user: {email}")

        # Authenticate user
        auth_response = supabase.auth.sign_in_with_password({
            "email": email,
            "password": password
        })

        if auth_response.user:
            # Get user data
            user_response = supabase.table('users').select('*').eq('id', auth_response.user.id).execute()

            if user_response.data:
                user_data = user_response.data[0]

                # Store user in session
                session.permanent = True
                session['user'] = {
                    'id': auth_response.user.id,
                    'email': auth_response.user.email,
                    'full_name': user_data.get('full_name', '')
                }

                response = jsonify({
                    "success": True,
                    "message": "Login successful",
                    "user": session['user']
                })
                response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
                response.headers.add('Access-Control-Allow-Credentials', 'true')
                return response, 200

    except Exception as e:
        error_message = handle_supabase_error(e)
        response = jsonify({
            "success": False,
            "message": error_message
        })
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response, 500

@app.route('/api/user', methods=['GET', 'OPTIONS'])
def get_user():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    if 'user' not in session:
        response = jsonify({
            "success": False,
            "message": "Not authenticated"
        })
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response, 401

    response = jsonify({
        "success": True,
        "user": session['user']
    })
    response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
    response.headers.add('Access-Control-Allow-Credentials', 'true')
    return response, 200

@app.route('/api/logout', methods=['POST', 'OPTIONS'])
def logout():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    session.clear()
    if supabase:
        supabase.auth.sign_out()

    response = jsonify({
        "success": True,
        "message": "Logout successful"
    })
    response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
    response.headers.add('Access-Control-Allow-Credentials', 'true')
    return response, 200

@app.route('/api/user/update', methods=['PUT', 'OPTIONS'])
def update_user_profile():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    if 'user' not in session:
        return jsonify({"success": False, "message": "Not authenticated"}), 401

    if not supabase:
        return jsonify({"success": False, "message": "Server configuration error"}), 500

    try:
        data = request.get_json()
        user_id = session['user']['id']

        # Fields that can be updated
        update_data = {}
        if 'full_name' in data:
            update_data['full_name'] = data['full_name']
        # Add other fields if you want to allow updating them (e.g., avatar_url)

        if not update_data:
            return jsonify({"success": False, "message": "No data provided for update"}), 400

        # Update user in 'users' table
        # Supabase doesn't directly update full_name via auth.update_user,
        # so we update our custom 'users' table.
        response = supabase.table('users').update(update_data).eq('id', user_id).execute()

        if response.data:
            # Update session with new data
            session['user']['full_name'] = update_data.get('full_name', session['user']['full_name'])

            return jsonify({
                "success": True,
                "message": "Profile updated successfully",
                "user": session['user']
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to update profile"}), 500

    except Exception as e:
        error_message = handle_supabase_error(e)
        return jsonify({"success": False, "message": error_message}), 500

@app.route('/api/user/change-password', methods=['POST', 'OPTIONS'])
def change_user_password():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "ok"})
        response.headers.add('Access-Control-Allow-Origin', 'http://localhost:5000')
        response.headers.add('Access-Control-Allow-Credentials', 'true')
        return response

    if 'user' not in session:
        return jsonify({"success": False, "message": "Not authenticated"}), 401

    if not supabase:
        return jsonify({"success": False, "message": "Server configuration error"}), 500

    try:
        data = request.get_json()
        current_password = data.get('current_password')
        new_password = data.get('new_password')

        if not current_password or not new_password:
            return jsonify({"success": False, "message": "Current and new passwords are required"}), 400

        # Supabase's auth.update_user does not take current_password for security reasons.
        # The user must be authenticated via a valid session.
        # If you want to verify current password, you'd need to re-authenticate the user
        # or rely on Supabase's session validity.
        # For simplicity, we'll assume the session is valid and directly update.
        # A more secure approach might involve a password reset flow or re-authentication.

        # Update password in Supabase Auth
        # Note: This requires the user's JWT to be valid and associated with the request.
        # The supabase.auth.update_user method typically works with the current session's user.
        auth_response = supabase.auth.update_user({"password": new_password})

        if auth_response.user:
            return jsonify({
                "success": True,
                "message": "Password changed successfully"
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to change password"}), 500

    except Exception as e:
        error_message = handle_supabase_error(e)
        return jsonify({"success": False, "message": error_message}), 500

@app.route('/api/generate-image', methods=['POST'])
def generate_image():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to generate images."}), 401

    global current_a4f_api_key_iterator # Declare global to modify the iterator
    global current_a4f_api_key # Declare global to modify the current key

    if not current_a4f_api_key_iterator:
        return jsonify({"success": False, "message": "Server configuration error: A4F API keys are not set."}), 500

    try:
        data = request.get_json()
        prompt = data.get('prompt')
        model = data.get('model', 'provider-4/imagen-4')
        size = data.get('size', '1024x1024')
        n = data.get('n', 1)

        if not prompt:
            return jsonify({"success": False, "message": "Image prompt is required."}), 400

        # Model-specific validation and parameter adjustment
        model_config = get_model_config(model)
        if not model_config:
            return jsonify({"success": False, "message": f"Unsupported model: {model}"}), 400

        # Validate parameters against model capabilities
        if n > model_config['max_images']:
            n = model_config['max_images']

        if size not in model_config['supported_sizes']:
            size = model_config['supported_sizes'][0]  # Default to first supported size

        # Prepare request with model-specific parameters
        payload = build_model_payload(model, prompt, n, size, model_config)

        a4f_url = "https://api.a4f.co/v1/images/generations"
        headers = {
            "Authorization": f"Bearer {current_a4f_api_key}", # Use the current key
            "Content-Type": "application/json"
        }

        print(f"Calling A4F.co API with model {model} and payload: {payload} using key: {current_a4f_api_key[:5]}...")

        # Attempt to make the request
        try:
            a4f_response = requests.post(a4f_url, headers=headers, data=json.dumps(payload), timeout=120)
            a4f_response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                print(f"Rate limit hit for key {current_a4f_api_key[:5]}... Switching to next key.")
                # Try to switch to the next key
                try:
                    current_a4f_api_key = next(current_a4f_api_key_iterator)
                    print(f"Switched to new key: {current_a4f_api_key[:5]}...")
                    # Re-attempt the request with the new key
                    headers["Authorization"] = f"Bearer {current_a4f_api_key}"
                    a4f_response = requests.post(a4f_url, headers=headers, data=json.dumps(payload), timeout=120)
                    a4f_response.raise_for_status()
                except StopIteration:
                    # This should not happen with itertools.cycle, but as a safeguard
                    print("No more A4F API keys available in the cycle.")
                    return jsonify({"success": False, "message": "All A4F API keys are rate-limited or exhausted."}), 503
                except requests.exceptions.HTTPError as retry_e:
                    # If the retry also fails with an HTTPError
                    error_detail = "AI service temporarily unavailable"
                    if retry_e.response.status_code == 429:
                        error_detail = "Rate limit exceeded on multiple keys. Please try again later."
                    elif retry_e.response.status_code == 402:
                        error_detail = "API quota exceeded. Please check your subscription."
                    return jsonify({"success": False, "message": error_detail}), retry_e.response.status_code
            else:
                # Re-raise other HTTP errors
                raise e

        a4f_data = a4f_response.json()

        # Parse response based on model
        image_urls = parse_model_response(a4f_data, model)

        if image_urls:
            return jsonify({
                "success": True,
                "message": f"Image(s) generated successfully using {model}.",
                "image_urls": image_urls,
                "model_used": model
            }), 200
        else:
            return jsonify({"success": False, "message": "AI API returned no image URLs."}), 500

    except requests.exceptions.HTTPError as e:
        error_detail = "AI service temporarily unavailable"
        if e.response.status_code == 429:
            error_detail = "Rate limit exceeded. Please try again later."
        elif e.response.status_code == 402:
            error_detail = "API quota exceeded. Please check your subscription."

        return jsonify({"success": False, "message": error_detail}), e.response.status_code
    except Exception as e:
        print(f"Error in generate_image: {e}")
        return jsonify({"success": False, "message": "An internal server error occurred."}), 500

@app.route('/api/chat/completions', methods=['POST'])
def chat_completions():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to use chat bots."}), 401

    global current_a4f_api_key_iterator
    global current_a4f_api_key

    if not current_a4f_api_key_iterator:
        return jsonify({"success": False, "message": "Server configuration error: A4F API keys are not set."}), 500

    try:
        data = request.get_json()
        model = data.get('model')
        messages = data.get('messages')
        temperature = data.get('temperature', 0.7)
        max_tokens = data.get('max_tokens', 500) # Default max_tokens for chat

        if not model or not messages:
            return jsonify({"success": False, "message": "Model and messages are required."}), 400

        # Validate model against supported chat models
        model_config = get_model_config(model)
        if not model_config or model_config.get('type') != 'chat':
            return jsonify({"success": False, "message": f"Unsupported or invalid chat model: {model}"}), 400

        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False # For now, we don't support streaming from backend to frontend
        }

        a4f_url = "https://api.a4f.co/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {current_a4f_api_key}",
            "Content-Type": "application/json"
        }

        print(f"Calling A4F.co Chat API with model {model} using key: {current_a4f_api_key[:5]}...")

        try:
            a4f_response = requests.post(a4f_url, headers=headers, data=json.dumps(payload), timeout=180) # Increased timeout for chat
            a4f_response.raise_for_status()
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                print(f"Rate limit hit for key {current_a4f_api_key[:5]}... Switching to next key.")
                try:
                    current_a4f_api_key = next(current_a4f_api_key_iterator)
                    print(f"Switched to new key: {current_a4f_api_key[:5]}...")
                    headers["Authorization"] = f"Bearer {current_a4f_api_key}"
                    a4f_response = requests.post(a4f_url, headers=headers, data=json.dumps(payload), timeout=180)
                    a4f_response.raise_for_status()
                except StopIteration:
                    print("No more A4F API keys available in the cycle.")
                    return jsonify({"success": False, "message": "All A4F API keys are rate-limited or exhausted."}), 503
                except requests.exceptions.HTTPError as retry_e:
                    error_detail = "AI service temporarily unavailable"
                    if retry_e.response.status_code == 429:
                        error_detail = "Rate limit exceeded on multiple keys. Please try again later."
                    elif retry_e.response.status_code == 402:
                        error_detail = "API quota exceeded. Please check your subscription."
                    return jsonify({"success": False, "message": error_detail}), retry_e.response.status_code
            else:
                raise e

        a4f_data = a4f_response.json()

        # Extract the assistant's message
        if a4f_data and 'choices' in a4f_data and len(a4f_data['choices']) > 0:
            assistant_message = a4f_data['choices'][0].get('message', {})
            return jsonify({
                "success": True,
                "message": "Chat completion successful.",
                "response": assistant_message,
                "usage": a4f_data.get('usage', {})
            }), 200
        else:
            return jsonify({"success": False, "message": "AI API returned no valid response."}), 500

    except requests.exceptions.HTTPError as e:
        error_detail = "AI service temporarily unavailable"
        if e.response.status_code == 429:
            error_detail = "Rate limit exceeded. Please try again later."
        elif e.response.status_code == 402:
            error_detail = "API quota exceeded. Please check your subscription."
        elif e.response.status_code == 400:
            error_detail = f"Bad request to AI service: {e.response.json().get('message', 'Invalid parameters')}"

        return jsonify({"success": False, "message": error_detail}), e.response.status_code
    except Exception as e:
        print(f"Error in chat_completions: {e}")
        return jsonify({"success": False, "message": "An internal server error occurred."}), 500

# FileName: MultipleFiles/app.py
# FileContents: ... (rest of the file)

@app.route('/api/enhance-image', methods=['POST'])
def enhance_image():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to enhance images."}), 401

    global current_picsart_api_key_iterator
    global current_picsart_api_key

    if not PICSART_API_KEYS: # Check the list directly, not the iterator
        return jsonify({"success": False, "message": "Server configuration error: Picsart API keys are not set."}), 500

    if 'image' not in request.files:
        return jsonify({"success": False, "message": "No image file provided."}), 400

    image_file = request.files['image']
    upscale_factor = request.form.get('upscale_factor', '2')
    output_format = request.form.get('format', 'JPG')

    picsart_url = "https://api.picsart.io/tools/1.0/upscale/enhance"

    # Prepare multipart/form-data
    files = {
        'image': (image_file.filename, image_file.stream, image_file.content_type)
    }
    data = {
        'upscale_factor': upscale_factor,
        'format': output_format
    }

    # Loop to try all available keys
    for _ in range(len(PICSART_API_KEYS)): # Iterate through all available keys
        headers = {
            "accept": "application/json",
            "X-Picsart-API-Key": current_picsart_api_key
        }

        print(f"Calling Picsart API with key: {current_picsart_api_key[:5]}...")

        try:
            picsart_response = requests.post(picsart_url, files=files, data=data, headers=headers, timeout=180)
            
            # Check for HTTP errors first
            picsart_response.raise_for_status() 
            
            # If no HTTP error, parse JSON and check for 'insufficient_credits'
            picsart_data = picsart_response.json()
            if picsart_data.get('status') == 'error' and 'insufficient_credits' in picsart_data.get('message', ''):
                print(f"Insufficient credits for Picsart key {current_picsart_api_key[:5]}... Switching to next key.")
                current_picsart_api_key = next(current_picsart_api_key_iterator)
                continue # Try next key
            
            # If successful or a different error, break the loop
            break 

        except requests.exceptions.HTTPError as e:
            error_message = e.response.json().get('message', 'Unknown error') if e.response.content else 'Unknown error'
            
            if e.response.status_code == 429:
                print(f"Rate limit hit for Picsart key {current_picsart_api_key[:5]}... Switching to next key.")
                current_picsart_api_key = next(current_picsart_api_key_iterator)
                continue # Try next key
            elif e.response.status_code == 402 or "insufficient_credits" in error_message: # Explicitly handle 402 and message
                print(f"Insufficient credits (402 or message) for Picsart key {current_picsart_api_key[:5]}... Switching to next key.")
                current_picsart_api_key = next(current_picsart_api_key_iterator)
                continue # Try next key
            else:
                # Re-raise other HTTP errors
                print(f"Picsart API Error: {e.response.status_code} - {e.response.text}")
                return jsonify({"success": False, "message": f"Picsart API error: {error_message}"}), e.response.status_code
        except StopIteration:
            # This should ideally not happen with the for loop, but as a safeguard
            print("No more Picsart API keys available in the cycle.")
            return jsonify({"success": False, "message": "All Picsart API keys are rate-limited or exhausted."}), 503
        except requests.exceptions.RequestException as e:
            print(f"Network or connection error during Picsart API call: {e}")
            return jsonify({"success": False, "message": "A network error occurred while connecting to the image enhancement service."}), 500
        except Exception as e:
            print(f"Error in enhance_image: {e}")
            return jsonify({"success": False, "message": "An internal server error occurred during image enhancement."}), 500

    # After the loop, check if a successful response was obtained
    if 'picsart_data' in locals() and picsart_data.get('status') == 'success' and 'data' in picsart_data and 'url' in picsart_data['data']:
        return jsonify({
            "success": True,
            "message": "Image enhanced successfully!",
            "enhanced_image_url": picsart_data['data']['url']
        }), 200
    else:
        # If loop finished without success, it means all keys failed or the last one failed in a non-rotatable way
        final_error_message = "All Picsart API keys are exhausted or invalid. Please check your API keys or try again later."
        if 'picsart_data' in locals() and picsart_data.get('message'):
            final_error_message = picsart_data.get('message')
        return jsonify({"success": False, "message": final_error_message}), 500



def get_model_config(model):
    """Get configuration for specific models - ONLY WORKING MODELS"""
    configs = {
        # Image Generation Models
        'provider-4/imagen-3': {
            'type': 'image',
            'max_images': 4,
            'supported_sizes': ['1024x1024', '1152x896', '896x1152'],
            'default_params': {}
        },
        'provider-4/imagen-4': {
            'type': 'image',
            'max_images': 4,
            'supported_sizes': ['1024x1024', '1152x896', '896x1152', '1216x832', '832x1216'],
            'default_params': {}
        },
        'provider-6/sana-1.5-flash': {
            'type': 'image',
            'max_images': 4,
            'supported_sizes': ['1024x1024', '1152x896', '896x1152'],
            'default_params': {}
        },
        'provider-6/sana-1.5': {
            'type': 'image',
            'max_images': 4,
            'supported_sizes': ['1024x1024', '1152x896', '896x1152'],
            'default_params': {}
        },
        'provider-3/FLUX.1-dev': {
            'type': 'image',
            'max_images': 2,
            'supported_sizes': ['1024x1024', '1152x896', '896x1152', '1344x768', '768x1344'],
            'default_params': {'guidance_scale': 7.5, 'num_inference_steps': 50}
        },
        # Chat Models
        'provider-6/gpt-4o-mini': {
            'type': 'chat',
            'description': 'GPT-4o mini: Fast, cost-effective, and capable of vision and function calling.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-6/gpt-5-nano': {
            'type': 'chat',
            'description': 'GPT-5 nano: A highly experimental, cutting-edge model with advanced capabilities.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-3/gpt-4': {
            'type': 'chat',
            'description': 'GPT-4: OpenAI\'s powerful general-purpose model, excellent for complex reasoning.',
            'features': [] # Cannot generate images
        },
        'provider-6/gpt-4.1-mini': {
            'type': 'chat',
            'description': 'GPT-4.1 mini: An optimized version of GPT-4 with vision and function calling.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-6/gpt-4.1': {
            'type': 'chat',
            'description': 'GPT-4.1: Enhanced version of GPT-4 with improved performance.',
            'features': []
        },
        'provider-6/o3-medium': {
            'type': 'chat',
            'description': 'O3 Medium: A versatile model from OpenAI, good for general tasks.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-3/deepseek-v3': {
            'type': 'chat',
            'description': 'Deepseek V3: Advanced model from Deepseek AI, strong in reasoning and multimodal tasks.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-6/deepseek-r1-0528-turbo': {
            'type': 'chat',
            'description': 'Deepseek R1 Turbo: Optimized for speed and reasoning tasks.',
            'features': ['Reasoning']
        },
        'provider-6/gemini-2.5-flash-thinking': {
            'type': 'chat',
            'description': 'Gemini 2.5 Flash Thinking: Google\'s fast and efficient model with strong reasoning.',
            'features': ['Vision', 'Reasoning']
        },
        'provider-6/gemini-2.5-flash': {
            'type': 'chat',
            'description': 'Gemini 2.5 Flash: Google\'s fastest model, ideal for quick responses and multimodal inputs.',
            'features': ['Vision', 'Function Calling']
        },
        'provider-2/gemini-2.0-flash': {
            'type': 'chat',
            'description': 'Gemini 2.0 Flash: A previous generation of Gemini Flash, still highly capable.',
            'features': ['Vision']
        },
        'provider-6/qwen-3-235b-a22b-2507': {
            'type': 'chat',
            'description': 'Qwen 3 235B A22B 2507: Alibaba\'s large language model, strong in reasoning and function calling.',
            'features': ['Reasoning', 'Function Calling']
        },
        'provider-3/qwen-3-235b-a22b': {
            'type': 'chat',
            'description': 'Qwen 3 235B A22B: A powerful Qwen model with hybrid reasoning capabilities.',
            'features': ['Hybrid-Reasoning', 'Function Calling']
        }
    }
    return configs.get(model)

def build_model_payload(model, prompt, n, size, model_config):
    """Build API payload with model-specific parameters"""
    payload = {
        "model": model,
        "prompt": prompt,
        "n": n,
        "size": size
    }

    # Add model-specific parameters
    if model_config.get('default_params'):
        payload.update(model_config['default_params'])

    # FLUX models might need special prompt formatting
    if 'FLUX' in model:
        payload['prompt'] = enhance_flux_prompt(prompt)

    return payload

def enhance_flux_prompt(prompt):
    """Enhance prompts for FLUX models"""
    # FLUX models respond well to detailed, descriptive prompts
    if len(prompt.split()) < 10:
        return f"{prompt}, highly detailed, professional quality, sharp focus"
    return prompt

def parse_model_response(response_data, model):
    """Parse API response to extract image URLs based on model"""
    # Default parsing logic for A4F API
    if response_data and 'data' in response_data and isinstance(response_data['data'], list):
        return [item['url'] for item in response_data['data'] if 'url' in item]

    # Fallback for different response structures
    if 'images' in response_data and isinstance(response_data['images'], list):
        return response_data['images']

    if 'urls' in response_data and isinstance(response_data['urls'], list):
        return response_data['urls']

    # Handle specific model response formats if needed
    if 'FLUX' in model and 'output' in response_data:
        # Handle FLUX-specific response format
        if isinstance(response_data['output'], list):
            return response_data['output']

    return []


# ... (rest of your existing app.py code) ...

@app.route('/api/compress-pdf', methods=['POST'])
def compress_pdf():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to compress PDFs."}), 401

    if 'pdf_file' not in request.files:
        return jsonify({"success": False, "message": "No PDF file provided."}), 400

    pdf_file = request.files['pdf_file']
    compression_level = request.form.get('compression_level', 'high')
    password = request.form.get('password')

    if pdf_file.filename == '':
        return jsonify({"success": False, "message": "No selected file."}), 400

    if not pdf_file.filename.lower().endswith('.pdf'):
        return jsonify({"success": False, "message": "Invalid file type. Only PDF files are allowed."}), 400

    original_filename = pdf_file.filename
    
    try:
        # Read the original file content
        pdf_content = pdf_file.read()
        original_size = len(pdf_content)
        
        print(f"Starting compression of {original_filename} ({original_size} bytes)")
        
        # Compress the PDF
        compressed_content = compress_pdf_advanced(pdf_content, compression_level)
        compressed_size = len(compressed_content)
        
        print(f"Compression result: {original_size} -> {compressed_size} bytes")
        
        # Add password protection if requested
        if password and password.strip():
            compressed_content = add_password_protection(compressed_content, password)
            compressed_size = len(compressed_content)
        
        # Create response
        compressed_filename = f"compressed_{original_filename}"
        
        response = make_response(send_file(
            io.BytesIO(compressed_content),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=compressed_filename
        ))
        
        # Add custom headers
        response.headers['X-Original-Size'] = str(original_size)
        response.headers['X-Compressed-Size'] = str(compressed_size)
        response.headers['X-Compressed-Filename'] = compressed_filename
        response.headers['Access-Control-Expose-Headers'] = 'X-Original-Size, X-Compressed-Size, X-Compressed-Filename'
        
        return response

    except Exception as e:
        print(f"Error during PDF compression: {e}")
        return jsonify({"success": False, "message": f"Compression failed: {str(e)}"}), 500


def compress_pdf_advanced(pdf_content, compression_level):
    """Advanced PDF compression using PyMuPDF with proper error handling"""
    try:
        # Open PDF from memory
        doc = fitz.open("pdf", pdf_content)
        
        # Set compression parameters based on level
        if compression_level == 'extreme':
            image_quality = 30
            max_image_size = (800, 1000)  # Max width, height
            remove_annotations = True
            remove_links = True
        elif compression_level == 'high':
            image_quality = 50
            max_image_size = (1200, 1500)
            remove_annotations = True
            remove_links = False
        elif compression_level == 'medium':
            image_quality = 70
            max_image_size = (1600, 2000)
            remove_annotations = False
            remove_links = False
        else:  # low
            image_quality = 85
            max_image_size = (2000, 2500)
            remove_annotations = False
            remove_links = False
        
        print(f"Processing {len(doc)} pages with {compression_level} compression")
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Remove annotations if requested
            if remove_annotations:
                try:
                    while page.first_annot:
                        page.delete_annot(page.first_annot)
                except:
                    pass
            
            # Remove links if requested
            if remove_links:
                try:
                    page.delete_link(page.first_link)
                except:
                    pass
            
            # Compress images on the page
            compress_page_images(page, image_quality, max_image_size)
        
        # Remove metadata for extreme compression
        if compression_level == 'extreme':
            try:
                doc.set_metadata({})
            except:
                pass
        
        # Save to memory with compression options
        output_buffer = io.BytesIO()
        
        # Use tobytes() method which is more reliable
        compressed_pdf = doc.tobytes(
            garbage=4,  # Remove unused objects
            deflate=True,  # Compress content streams
            clean=True,  # Clean up the file structure
            ascii=False,  # Use binary encoding
            expand=0,  # Don't expand content streams
            linear=False,  # Don't linearize
            pretty=False  # Don't pretty-print
        )
        
        doc.close()
        
        print(f"PyMuPDF compression successful: {len(pdf_content)} -> {len(compressed_pdf)} bytes")
        return compressed_pdf
        
    except Exception as e:
        print(f"PyMuPDF compression failed: {e}")
        # Fallback to pikepdf compression
        return compress_pdf_fallback(pdf_content, compression_level)


def compress_page_images(page, quality, max_size):
    """Compress images on a PDF page"""
    try:
        # Get all images on the page
        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            try:
                # Get image reference
                xref = img[0]
                
                # Extract image
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Skip very small images
                if len(image_bytes) < 5000:  # Less than 5KB
                    continue
                
                print(f"Processing image {img_index}: {len(image_bytes)} bytes, format: {image_ext}")
                
                # Open image with PIL
                try:
                    image = Image.open(io.BytesIO(image_bytes))
                except Exception as e:
                    print(f"Could not open image {img_index}: {e}")
                    continue
                
                # Get original dimensions
                width, height = image.size
                max_width, max_height = max_size
                
                # Resize if too large
                if width > max_width or height > max_height:
                    # Calculate resize ratio
                    ratio = min(max_width / width, max_height / height)
                    new_width = int(width * ratio)
                    new_height = int(height * ratio)
                    
                    print(f"Resizing image from {width}x{height} to {new_width}x{new_height}")
                    image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to RGB if necessary
                if image.mode in ('RGBA', 'LA', 'P'):
                    # Create white background
                    background = Image.new('RGB', image.size, (255, 255, 255))
                    if image.mode == 'P':
                        image = image.convert('RGBA')
                    if image.mode in ('RGBA', 'LA'):
                        background.paste(image, mask=image.split()[-1])
                    else:
                        background.paste(image)
                    image = background
                elif image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Compress image to JPEG
                compressed_image_buffer = io.BytesIO()
                image.save(
                    compressed_image_buffer, 
                    format='JPEG', 
                    quality=quality, 
                    optimize=True,
                    progressive=True
                )
                compressed_image_bytes = compressed_image_buffer.getvalue()
                
                # Only replace if compression was effective
                if len(compressed_image_bytes) < len(image_bytes):
                    print(f"Image compressed: {len(image_bytes)} -> {len(compressed_image_bytes)} bytes")
                    
                    # Replace image in PDF
                    try:
                        # Create new image object
                        img_dict = {
                            "width": image.size[0],
                            "height": image.size[1],
                            "colorspace": 3,  # RGB
                            "bpc": 8,  # 8 bits per component
                            "filter": "DCTDecode",  # JPEG compression
                            "image": compressed_image_bytes
                        }
                        
                        # Update the image
                        page.parent.update_stream(xref, compressed_image_bytes)
                        
                    except Exception as e:
                        print(f"Could not replace image {img_index}: {e}")
                else:
                    print(f"Image {img_index}: compression not effective, keeping original")
                
            except Exception as e:
                print(f"Error processing image {img_index}: {e}")
                continue
                
    except Exception as e:
        print(f"Page image compression error: {e}")
        pass


def compress_pdf_fallback(pdf_content, compression_level):
    """Fallback compression using pikepdf only"""
    try:
        print("Using pikepdf fallback compression")
        
        # Open PDF from memory
        input_stream = io.BytesIO(pdf_content)
        output_stream = io.BytesIO()
        
        with pikepdf.Pdf.open(input_stream) as pdf:
            # Remove unused resources
            try:
                pdf.remove_unreferenced_resources()
            except:
                pass
            
            # Remove metadata for aggressive compression
            if compression_level in ['extreme', 'high']:
                try:
                    # Clear document info
                    if hasattr(pdf, 'docinfo') and pdf.docinfo:
                        for key in list(pdf.docinfo.keys()):
                            try:
                                del pdf.docinfo[key]
                            except:
                                pass
                    
                    # Remove XMP metadata
                    if '/Metadata' in pdf.Root:
                        del pdf.Root['/Metadata']
                except:
                    pass
            
            # Save with compression
            pdf.save(
                output_stream,
                compress_streams=True,
                object_stream_mode=pikepdf.ObjectStreamMode.generate,
                normalize_content=True
            )
        
        compressed_content = output_stream.getvalue()
        print(f"Pikepdf compression: {len(pdf_content)} -> {len(compressed_content)} bytes")
        return compressed_content
        
    except Exception as e:
        print(f"Fallback compression failed: {e}")
        # Return original content if all compression fails
        return pdf_content


def add_password_protection(pdf_content, password):
    """Add password protection to PDF content"""
    try:
        input_stream = io.BytesIO(pdf_content)
        output_stream = io.BytesIO()
        
        with pikepdf.Pdf.open(input_stream) as pdf:
            encryption = pikepdf.Encryption(
                owner=password,
                user=password,
                R=6,
                allow=pikepdf.Permissions(
                    accessibility=True,
                    extract=True,
                    modify_annotation=True,
                    modify_assembly=False,
                    modify_form=True,
                    modify_other=False,
                    print_highres=True,
                    print_lowres=True
                )
            )
            pdf.save(output_stream, encryption=encryption)
        
        return output_stream.getvalue()
        
    except Exception as e:
        print(f"Password protection failed: {e}")
        return pdf_content

# --- Unit Converter Data and Logic ---
UNIT_CATEGORIES = {
    "length": {
        "meter": 1, "kilometer": 1000, "centimeter": 0.01, "millimeter": 0.001,
        "micrometer": 1e-6, "nanometer": 1e-9, "mile": 1609.34, "yard": 0.9144,
        "foot": 0.3048, "inch": 0.0254, "light_year": 9.461e15, "nautical_mile": 1852
    },
    "mass": {
        "kilogram": 1, "gram": 0.001, "milligram": 1e-6, "microgram": 1e-9,
        "tonne": 1000, "pound": 0.453592, "ounce": 0.0283495, "stone": 6.35029
    },
    "volume": {
        "liter": 1, "milliliter": 0.001, "cubic_meter": 1000, "cubic_centimeter": 0.001,
        "gallon_us": 3.78541, "quart_us": 0.946353, "pint_us": 0.473176,
        "fluid_ounce_us": 0.0295735, "cup_us": 0.236588
    },
    "temperature": {
        "celsius": "C", "fahrenheit": "F", "kelvin": "K"
    },
    "time": {
        "second": 1, "millisecond": 0.001, "minute": 60, "hour": 3600,
        "day": 86400, "week": 604800, "month": 2.628e6, "year": 3.154e7
    },
    "speed": {
        "meter_per_second": 1, "kilometer_per_hour": 0.277778, "mile_per_hour": 0.44704,
        "knot": 0.514444, "foot_per_second": 0.3048
    },
    "area": {
        "square_meter": 1, "square_kilometer": 1e6, "square_centimeter": 1e-4,
        "square_millimeter": 1e-6, "acre": 4046.86, "hectare": 10000,
        "square_mile": 2.59e6, "square_yard": 0.836127, "square_foot": 0.092903,
        "square_inch": 0.00064516
    },
    "pressure": {
        "pascal": 1, "kilopascal": 1000, "bar": 100000, "psi": 6894.76,
        "atmosphere": 101325, "torr": 133.322
    },
    "energy": {
        "joule": 1, "kilojoule": 1000, "calorie": 4.184, "kilocalorie": 4184,
        "watt_hour": 3600, "kilowatt_hour": 3.6e6, "electronvolt": 1.60218e-19
    },
    "data_storage": {
        "bit": 1, "byte": 8, "kilobit": 1000, "kilobyte": 8000,
        "megabit": 1e6, "megabyte": 8e6, "gigabit": 1e9, "gigabyte": 8e9,
        "terabit": 1e12, "terabyte": 8e12
    }
}

def convert_unit(value, from_unit, to_unit, category):
    if category not in UNIT_CATEGORIES:
        raise ValueError("Invalid category")

    units = UNIT_CATEGORIES[category]

    if from_unit not in units or to_unit not in units:
        raise ValueError("Invalid units for the selected category")

    if category == "temperature":
        # Temperature conversions are special
        if from_unit == "celsius" and to_unit == "fahrenheit":
            return (value * 9/5) + 32
        elif from_unit == "fahrenheit" and to_unit == "celsius":
            return (value - 32) * 5/9
        elif from_unit == "celsius" and to_unit == "kelvin":
            return value + 273.15
        elif from_unit == "kelvin" and to_unit == "celsius":
            return value - 273.15
        elif from_unit == "fahrenheit" and to_unit == "kelvin":
            return (value - 32) * 5/9 + 273.15
        elif from_unit == "kelvin" and to_unit == "fahrenheit":
            return (value - 273.15) * 9/5 + 32
        else:
            return value # Same unit, or unsupported temperature conversion
    else:
        # Convert 'from_unit' to base unit (e.g., meter, kilogram)
        value_in_base = value * units[from_unit]
        # Convert base unit to 'to_unit'
        converted_value = value_in_base / units[to_unit]
        return converted_value

@app.route('/api/unit-converter/categories', methods=['GET'])
def get_unit_categories():
    return jsonify({
        "success": True,
        "categories": list(UNIT_CATEGORIES.keys())
    })

@app.route('/api/unit-converter/units/<category>', methods=['GET'])
def get_units_by_category(category):
    if category not in UNIT_CATEGORIES:
        return jsonify({"success": False, "message": "Invalid category"}), 400
    return jsonify({
        "success": True,
        "units": list(UNIT_CATEGORIES[category].keys())
    })

@app.route('/api/unit-converter/convert', methods=['POST'])
def unit_convert():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to use the unit converter."}), 401

    data = request.get_json()
    value = data.get('value')
    from_unit = data.get('from_unit')
    to_unit = data.get('to_unit')
    category = data.get('category')

    if not all([value is not None, from_unit, to_unit, category]):
        return jsonify({"success": False, "message": "Missing conversion parameters."}), 400

    try:
        value = float(value)
    except ValueError:
        return jsonify({"success": False, "message": "Value must be a number."}), 400

    try:
        converted_value = convert_unit(value, from_unit, to_unit, category)
        return jsonify({
            "success": True,
            "original_value": value,
            "from_unit": from_unit,
            "to_unit": to_unit,
            "converted_value": converted_value,
            "category": category
        })
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400
    except Exception as e:
        print(f"Error during unit conversion: {e}")
        return jsonify({"success": False, "message": "An internal server error occurred during conversion."}), 500


# ... (rest of your existing app.py code) ...

@app.route('/api/convert-images-to-pdf', methods=['POST'])
def convert_images_to_pdf():
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to convert images to PDF."}), 401

    if 'images' not in request.files:
        return jsonify({"success": False, "message": "No image files provided."}), 400

    images = request.files.getlist('images')
    if not images:
        return jsonify({"success": False, "message": "No image files provided."}), 400

    page_size_str = request.form.get('pageSize', 'A4')
    page_orientation = request.form.get('pageOrientation', 'portrait')
    image_fit = request.form.get('imageFit', 'fit')
    jpeg_quality = int(request.form.get('jpegQuality', 90))
    password = request.form.get('password')

    custom_width = None
    custom_height = None
    if page_size_str == 'Custom':
        try:
            custom_width = float(request.form.get('customWidth'))
            custom_height = float(request.form.get('customHeight'))
        except (ValueError, TypeError):
            return jsonify({"success": False, "message": "Invalid custom page size dimensions."}), 400

    # Determine page size
    page_size = A4 # Default
    if page_size_str == 'Letter':
        page_size = LETTER
    elif page_size_str == 'Legal':
        page_size = LEGAL
    elif page_size_str == 'A3':
        page_size = A3
    elif page_size_str == 'Custom' and custom_width and custom_height:
        page_size = (custom_width, custom_height) # Custom size in points (px)
    # A4 is default if not matched

    # Apply orientation
    if page_orientation == 'landscape':
        page_size = (page_size[1], page_size[0]) # Swap width and height for landscape

    output_buffer = io.BytesIO()
    c = canvas.Canvas(output_buffer, pagesize=page_size)

    try:
        for i, image_file in enumerate(images):
            try:
                img_stream = io.BytesIO(image_file.read())
                img = Image.open(img_stream)

                # Convert image to RGB if it's not (e.g., RGBA, P)
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create a white background image
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA') # Convert paletted to RGBA first
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')

                # Save as JPEG to a new buffer to control quality and format
                temp_img_buffer = io.BytesIO()
                img.save(temp_img_buffer, format='JPEG', quality=jpeg_quality)
                temp_img_buffer.seek(0) # Rewind buffer

                # Create ReportLab ImageReader
                reportlab_img = ImageReader(temp_img_buffer)

                # Calculate image position and size on PDF page
                img_width, img_height = reportlab_img.getSize()
                page_width, page_height = page_size

                if image_fit == 'fit':
                    # Scale image to fit within page, maintaining aspect ratio
                    aspect_ratio = img_width / img_height
                    if page_width / page_height > aspect_ratio:
                        # Page is wider than image, fit by height
                        draw_height = page_height
                        draw_width = page_height * aspect_ratio
                    else:
                        # Page is taller than image, fit by width
                        draw_width = page_width
                        draw_height = page_width / aspect_ratio
                elif image_fit == 'fill':
                    # Scale image to fill page, potentially cropping
                    aspect_ratio = img_width / img_height
                    if page_width / page_height < aspect_ratio:
                        # Page is taller than image, fill by height (will crop width)
                        draw_height = page_height
                        draw_width = page_height * aspect_ratio
                    else:
                        # Page is wider than image, fill by width (will crop height)
                        draw_width = page_width
                        draw_height = page_width / aspect_ratio
                else: # original
                    draw_width = img_width
                    draw_height = img_height

                # Center the image on the page
                x = (page_width - draw_width) / 2
                y = (page_height - draw_height) / 2

                c.drawImage(reportlab_img, x, y, width=draw_width, height=draw_height)
                
                if i < len(images) - 1: # Add new page for all but the last image
                    c.showPage()

            except Exception as e:
                print(f"Error processing image {image_file.filename}: {e}")
                # Decide how to handle errors: skip image, return error, etc.
                # For now, we'll just log and continue, but a more robust solution might skip or fail.
                continue

        c.save()
        output_buffer.seek(0)

        final_pdf_content = output_buffer.getvalue()

        # Add password protection if requested
        if password and password.strip():
            final_pdf_content = add_password_protection_reportlab(final_pdf_content, password)

        response = make_response(send_file(
            io.BytesIO(final_pdf_content),
            mimetype='application/pdf',
            as_attachment=True,
            download_name='converted_images.pdf'
        ))
        response.headers['X-Converted-Filename'] = 'converted_images.pdf' # Custom header for filename
        response.headers['Access-Control-Expose-Headers'] = 'X-Converted-Filename' # Expose custom header

        return response

    except Exception as e:
        print(f"Error during image to PDF conversion: {e}")
        return jsonify({"success": False, "message": f"Conversion failed: {str(e)}"}), 500

def add_password_protection_reportlab(pdf_content, password):
    """Add password protection to PDF content using pikepdf (more robust than ReportLab's built-in)"""
    try:
        input_stream = io.BytesIO(pdf_content)
        output_stream = io.BytesIO()

        with pikepdf.Pdf.open(input_stream) as pdf:
            encryption = pikepdf.Encryption(
                owner=password,
                user=password,
                R=6, # AES-256 encryption
                allow=pikepdf.Permissions(
                    accessibility=True,
                    extract=True,
                    modify_annotation=True,
                    modify_assembly=False,
                    modify_form=True,
                    modify_other=False,
                    print_highres=True,
                    print_lowres=True
                )
            )
            pdf.save(output_stream, encryption=encryption)

        return output_stream.getvalue()

    except Exception as e:
        print(f"Password protection failed for image-to-pdf: {e}")
        # If password protection fails, return the unencrypted PDF
        return pdf_content

# ... (rest of your existing app.py code) ...


@app.route('/api/convert-pdf-to-images', methods=['POST'])
def convert_pdf_to_images():
    # Authentication check
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to convert PDF to images."}), 401

    # Validate PDF file
    if 'pdf' not in request.files:
        return jsonify({"success": False, "message": "No PDF file provided."}), 400

    pdf_file = request.files['pdf']
    if not pdf_file or pdf_file.filename == '':
        return jsonify({"success": False, "message": "No PDF file provided."}), 400

    # Get form parameters
    image_format = request.form.get('imageFormat', 'png').lower()
    image_quality = int(request.form.get('imageQuality', 90))
    dpi = int(request.form.get('dpi', 150))
    download_option = request.form.get('downloadOption', 'zip')

    # Validate parameters
    if image_format not in ['png', 'jpeg']:
        return jsonify({"success": False, "message": "Invalid image format. Use 'png' or 'jpeg'."}), 400
    if not (1 <= image_quality <= 100):
        return jsonify({"success": False, "message": "Image quality must be between 1 and 100."}), 400
    if dpi not in [72, 150, 300]:
        return jsonify({"success": False, "message": "Invalid DPI. Use 72, 150, or 300."}), 400
    if download_option not in ['zip', 'individual']:
        return jsonify({"success": False, "message": "Invalid download option. Use 'zip' or 'individual'."}), 400

    try:
        # Read PDF file into memory
        pdf_content = pdf_file.read()

        # Open PDF with PyMuPDF
        pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
        page_count = pdf_document.page_count

        if page_count == 0:
            pdf_document.close()
            return jsonify({"success": False, "message": "No pages found in the PDF."}), 400

        images = []
        # Convert each page to an image
        for page_num in range(page_count):
            page = pdf_document.load_page(page_num)
            # Set zoom factor based on DPI (1 pixel = 1/72 inch at 72 DPI)
            zoom = dpi / 72
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix)
            
            # Convert PyMuPDF pixmap to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)

        pdf_document.close()

        if download_option == 'zip':
            # Create ZIP file in memory
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for i, img in enumerate(images):
                    # Convert image to RGB if needed (e.g., for JPEG)
                    if image_format == 'jpeg' and img.mode != 'RGB':
                        img = img.convert('RGB')

                    # Save image to buffer
                    img_buffer = io.BytesIO()
                    if image_format == 'jpeg':
                        img.save(img_buffer, format='JPEG', quality=image_quality)
                    else:
                        img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Write to ZIP
                    zip_file.writestr(f'page_{i + 1}.{image_format}', img_buffer.getvalue())

            zip_buffer.seek(0)

            # Prepare response
            response = make_response(send_file(
                zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name='converted_images.zip'
            ))
            response.headers['X-Converted-Filename'] = 'converted_images.zip'
            response.headers['X-Page-Count'] = str(len(images))
            response.headers['Access-Control-Expose-Headers'] = 'X-Converted-Filename,X-Page-Count'
            return response

        else:  # download_option == 'individual'
            # Return the first page as an image (client expects single blob)
            img = images[0]
            if image_format == 'jpeg' and img.mode != 'RGB':
                img = img.convert('RGB')

            img_buffer = io.BytesIO()
            if image_format == 'jpeg':
                img.save(img_buffer, format='JPEG', quality=image_quality)
            else:
                img.save(img_buffer, format='PNG')
            img_buffer.seek(0)

            response = make_response(send_file(
                img_buffer,
                mimetype=f'image/{image_format}',
                as_attachment=True,
                download_name=f'page_1.{image_format}'
            ))
            response.headers['X-Converted-Filename'] = f'page_1.{image_format}'
            response.headers['X-Page-Count'] = str(len(images))
            response.headers['Access-Control-Expose-Headers'] = 'X-Converted-Filename,X-Page-Count'
            return response

    except Exception as e:
        print(f"Error during PDF to image conversion: {e}")
        return jsonify({"success": False, "message": f"Conversion failed: {str(e)}"}), 500

@app.route('/api/convert-pdf-to-word', methods=['POST'])
def convert_pdf_to_word():
    # Authentication check
    if 'user' not in session:
        return jsonify({"success": False, "message": "Authentication required to convert PDF to Word."}), 401

    # Validate PDF file
    if 'pdf' not in request.files:
        return jsonify({"success": False, "message": "No PDF file provided."}), 400

    pdf_file = request.files['pdf']
    if not pdf_file or pdf_file.filename == '':
        return jsonify({"success": False, "message": "No PDF file provided."}), 400

    # Get form parameters
    layout_preservation = request.form.get('layoutPreservation', 'full').lower()

    # Validate parameters
    if layout_preservation not in ['full', 'text']:
        return jsonify({"success": False, "message": "Invalid layout preservation option. Use 'full' or 'text'."}), 400

    try:
        # Read PDF file into memory
        pdf_content = pdf_file.read()

        # Open PDF with PyMuPDF for page count and image rendering
        pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
        page_count = pdf_document.page_count

        if page_count == 0:
            pdf_document.close()
            return jsonify({"success": False, "message": "No pages found in the PDF."}), 400

        # Initialize Word document
        doc = Document()

        # Try extracting text for text-based PDFs
        pdf_reader = PdfReader(io.BytesIO(pdf_content))
        text_content = ""
        has_text = False

        # Check if PDF has extractable text
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                has_text = True
                text_content += text + "\n\n"

        if layout_preservation == 'full':
            # Convert each page to an image and add text (if available)
            for page_num in range(page_count):
                page = pdf_document.load_page(page_num)
                zoom = 300 / 72  # 300 DPI for high-quality images
                matrix = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=matrix)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                doc.add_picture(img_buffer, width=Inches(6.0))  # Adjust width as needed
                if has_text:
                    text = pdf_reader.pages[page_num].extract_text()
                    if text and text.strip():
                        doc.add_paragraph(text)
                if page_num < page_count - 1:
                    doc.add_page_break()
        else:
            # Text-only mode
            if not has_text:
                pdf_document.close()
                return jsonify({"success": False, "message": "No extractable text found in the PDF for text-only mode."}), 400
            doc.add_paragraph(text_content)

        pdf_document.close()

        # Save Word document to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Prepare response
        output_filename = secure_filename(os.path.splitext(pdf_file.filename)[0] + '.docx')
        response = make_response(send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        ))
        response.headers['X-Converted-Filename'] = output_filename
        response.headers['X-Page-Count'] = str(page_count)
        response.headers['Access-Control-Expose-Headers'] = 'X-Converted-Filename,X-Page-Count'
        return response

    except Exception as e:
        print(f"Error during PDF to Word conversion: {e}")
        return jsonify({"success": False, "message": f"Conversion failed: {str(e)}"}), 500


if __name__ == '__main__':
    print("=" * 50)
    print("Starting Flask Server...")
    print("CORS configured for all origins")
    print("=" * 50)
   

